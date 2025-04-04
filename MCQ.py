import streamlit as st
import openai
import google.generativeai as genai

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

import os
from dotenv import load_dotenv

# Load API keys from .env
load_dotenv()

# Access the keys
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# Access the keys
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)

openai_key = OPENAI_API_KEY


# Function to get MCQ questions from OpenAI GPT-3.5
def generate_mcq_questions(topic, difficulty, num_questions):
    prompt =  f"""
    Generate a multiple-choice quiz with the following specifications:
    - Topic: "{topic}"
    - Difficulty level: "{difficulty}"
    - Number of questions: {num_questions}
    - Each question should have 4 options labeled a, b, c, and d.
    - Each option should be brief (2-3 words).
    - Clearly specify the correct answer for each question.

    Ensure the questions and options are:
    1. Clear and concise.
    2. Relevant to the specified topic.
    3. Appropriately challenging based on the specified difficulty level.
    4. Complete with all parts included.

    Do not include code-related questions.

    Example format:
    Q1: What is the capital of France?
    a. Berlin
    b. Madrid
    c. Paris
    d. Rome
    Answer: c

    Quiz:
    """
    
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)
    
    return response.text
# Function to format the generated quiz to exclude answers and add selection options
def format_quiz(quiz):
    lines = quiz.split("\n")
    formatted_quiz = []
    current_question = []
    for line in lines:
        if re.match(r"^Q\d+: ", line):
            if current_question:
                formatted_quiz.append(current_question)
            current_question = [line]
        elif re.match(r"^[a-d]\. ", line):
            current_question.append(line)
        elif re.match(r"^Answer: ", line):
            current_question.append(line)
    if current_question:
        formatted_quiz.append(current_question)
    return formatted_quiz

# Function to generate a DOCX document
def generate_docx(quiz, heading1, heading2):
    doc = Document()
    
    heading1_paragraph = doc.add_heading(heading1, level=0)
    heading1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    heading2_paragraph = doc.add_heading(heading2, level=2)
    heading2_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Name:")
    doc.add_paragraph("Roll number:")
    doc.add_paragraph("Class:")
    doc.add_paragraph("Section:")
    doc.add_paragraph("")  # Add space after the fields

    for question in quiz:
        for line in question:
            if not re.match(r"^Answer: ", line):
                doc.add_paragraph(line)
        doc.add_paragraph("")  # Add a line break between questions
    
    # Save document to BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

# Streamlit app
def MCQ():
    st.title("MCQ Quiz Generator")

    # Input fields for quiz parameters
    institute_name = st.text_input("Enter the Name of Institute:")
    quiz_title = st.text_input("Enter the quiz title:")
    topic = st.text_input("Enter the topic:")
    difficulty = st.selectbox("Select the difficulty level:", ["Beginner", "Intermediate", "Expert"])
    num_questions = st.number_input("Enter the number of questions:", min_value=1, max_value=20, value=5)

    # Generate quiz button
    if st.button("Generate Quiz"):
        if topic:
            with st.spinner("Generating questions..."):
                quiz = generate_mcq_questions(topic, difficulty, num_questions)
                formatted_quiz = format_quiz(quiz)
                st.subheader("Generated Quiz:")
                if institute_name:
                    st.write(f"{institute_name}")
                if quiz_title:
                    st.write(f"{quiz_title}")
                
                # Store the quiz in session state
                st.session_state['quiz'] = formatted_quiz
                
                # Generate DOCX and provide download button
                docx_content = generate_docx(formatted_quiz, institute_name, quiz_title)
                st.session_state['docx_content'] = docx_content
        else:
            st.error("Please enter a topic to generate the quiz.")

    # Display the quiz and allow users to view options
    if 'quiz' in st.session_state:
        for i, question in enumerate(st.session_state['quiz']):
            st.write(question[0])
            options = question[1:]
            for line in options:
                if "Answer: " in line:
                    st.write(f"Correct answer: {line.split(': ')[1]}")
                else:
                    st.write(line)

        # Download DOCX button
        if 'quiz' in st.session_state and 'docx_content' in st.session_state:
            st.download_button(
                label="Download Quiz as DOCX",
                data=st.session_state['docx_content'],
                file_name=f"{topic} quiz.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )